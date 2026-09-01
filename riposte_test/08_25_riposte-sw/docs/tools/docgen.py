#!/usr/bin/env python3
"""문서 블록 모델 + Markdown/Word 렌더러.

`make_*.py` 는 문서 내용을 블록 리스트로만 정의하고, 렌더링은 여기서 한다.
같은 내용을 .md 와 .docx 에 각각 적어 두면 반드시 어긋나기 때문이다 — 내용은
한 곳에만 있고, 출력 형식이 둘인 구조를 강제한다.

블록:
    ("h", level, text)          제목 (level 0 = 문서 제목)
    ("p", text)                 문단
    ("p_bold", text)            강조 문단
    ("p_note", text)            보조 설명 (기울임 / md 는 인용)
    ("ul", [text, ...])         글머리 목록
    ("table", [헤더...], [[행...]], [폭...])   표 (폭은 docx 전용, cm)
    ("code", text)              고정폭 블록 (다이어그램)

.docx 는 python-docx 가 있을 때만 만든다. 없으면 .md 만 쓰고 그 사실을 알린다.
"""
import os

FONT = "맑은 고딕"
MONO = "D2Coding"


# --------------------------------------------------------------- Markdown
def render_markdown(blocks, path):
    out = []
    for b in blocks:
        kind = b[0]
        if kind == "h":
            level, text = b[1], b[2]
            out.append(("#" * max(1, level + 1)) + " " + text.replace("\n", " — "))
            out.append("")
        elif kind == "p":
            out.append(b[1])
            out.append("")
        elif kind == "p_bold":
            out.append("**" + b[1] + "**")
            out.append("")
        elif kind == "p_note":
            out.append("> " + b[1])
            out.append("")
        elif kind == "ul":
            for item in b[1]:
                out.append("- " + item)
            out.append("")
        elif kind == "code":
            out.append("```")
            out.append(b[1])
            out.append("```")
            out.append("")
        elif kind == "table":
            headers, rows = b[1], b[2]
            out.append("| " + " | ".join(headers) + " |")
            out.append("|" + "|".join(["---"] * len(headers)) + "|")
            for row in rows:
                cells = [str(c).replace("|", "\\|").replace("\n", " ") for c in row]
                out.append("| " + " | ".join(cells) + " |")
            out.append("")
        else:
            raise ValueError("unknown block: " + kind)
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(out).rstrip() + "\n")
    return path


# --------------------------------------------------------------- Word
def render_docx(blocks, path):
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError:
        return None

    accent = RGBColor(0x1F, 0x3A, 0x5F)

    def kor(run, name=FONT):
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT)
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    for b in blocks:
        kind = b[0]
        if kind == "h":
            level, text = b[1], b[2]
            p = doc.add_heading(text, level=level)
            if level == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = accent
                kor(run)
        elif kind in ("p", "p_bold", "p_note"):
            p = doc.add_paragraph()
            run = p.add_run(b[1])
            run.bold = kind == "p_bold"
            run.italic = kind == "p_note"
            run.font.size = Pt(9 if kind == "p_note" else 10)
            kor(run)
        elif kind == "ul":
            for item in b[1]:
                p = doc.add_paragraph(item, style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(10)
                    kor(run)
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(b[1])
            run.font.size = Pt(8)
            kor(run, MONO)
        elif kind == "table":
            headers, rows = b[1], b[2]
            widths = b[3] if len(b) > 3 else None
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, name in enumerate(headers):
                cell = t.rows[0].cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(name)
                run.bold = True
                run.font.size = Pt(9)
                kor(run)
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ""
                    run = cells[i].paragraphs[0].add_run(str(val))
                    run.font.size = Pt(9)
                    kor(run)
            if widths:
                for r in t.rows:
                    for i, w in enumerate(widths):
                        r.cells[i].width = Cm(w)
            doc.add_paragraph()

    doc.save(path)
    return path


def emit(blocks, md_path, docx_path):
    """Both renderings from one content definition."""
    render_markdown(blocks, md_path)
    print("saved:", md_path)
    made = render_docx(blocks, docx_path)
    if made:
        print("saved:", made)
    else:
        print("SKIP docx: python-docx not installed (pip install python-docx);",
              "the Markdown above is complete on its own")
    return made is not None


def check_parallel(blocks_a, blocks_b):
    """Two language versions must have the same SHAPE.

    Translations drift silently: a section added to one language and forgotten
    in the other produces two documents that disagree while both look complete.
    This cannot verify meaning, but it does catch the structural half — a
    missing section, a table that gained a row, a list that lost an item — and
    it fails at generation time rather than in review.
    """
    if len(blocks_a) != len(blocks_b):
        raise AssertionError(
            f"block count differs: {len(blocks_a)} vs {len(blocks_b)} — "
            "a section was added or removed in one language only")
    for i, (a, b) in enumerate(zip(blocks_a, blocks_b)):
        if a[0] != b[0]:
            raise AssertionError(f"block {i}: kind {a[0]} vs {b[0]}")
        if a[0] == "h" and a[1] != b[1]:
            raise AssertionError(f"block {i}: heading level {a[1]} vs {b[1]}")
        if a[0] == "ul" and len(a[1]) != len(b[1]):
            raise AssertionError(
                f"block {i}: list has {len(a[1])} vs {len(b[1])} items")
        if a[0] == "table":
            if len(a[1]) != len(b[1]):
                raise AssertionError(f"block {i}: table column count differs")
            if len(a[2]) != len(b[2]):
                raise AssertionError(
                    f"block {i}: table has {len(a[2])} vs {len(b[2])} rows")


def out_dir():
    """docs/ — this file lives in docs/tools/."""
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
